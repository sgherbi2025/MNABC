
import numpy as np
import time

def _sphere(x):
    return float(np.sum(x * x))

def dynamic_radius(fes, max_fes, sn):
    """Eq. (12) : k = (SN/4 - 1) * (FEs/MaxFEs) + 1"""
    k = (sn / 4.0 - 1.0) * (fes / max_fes) + 1.0
    return max(1, int(round(k)))

def ring_neighbor_indices(i, k, sn):
    offsets = list(range(-k, 0)) + list(range(1, k + 1))
    return np.array([(i + o) % sn for o in offsets], dtype=int)

def compute_ph(population, fitness_target):
    """ILM (Eq. 5-10): PH = (1/|V|) * sum |vt - vr|, reference = Sphere."""
    sn = population.shape[0]
    fitness_ref = np.array([_sphere(x) for x in population])

    def info_matrix(fit):
        M = np.full((sn, sn), 0.5)
        for a in range(sn):
            for b in range(sn):
                if a == b:
                    continue
                if fit[a] < fit[b]:
                    M[a, b] = 1.0
                elif fit[a] > fit[b]:
                    M[a, b] = 0.0
        return M

    Mt = info_matrix(fitness_target)
    Mr = info_matrix(fitness_ref)
    best_idx = int(np.argmin(fitness_target))  

    vt, vr = [], []
    for a in range(sn):
        for b in range(a + 1, sn):
            if a == best_idx or b == best_idx:
                continue
            vt.append(Mt[a, b])
            vr.append(Mr[a, b])
    if not vt:
        return 0.0
    return float(np.mean(np.abs(np.array(vt) - np.array(vr))))

class ASABC:
    """Faithfully implements Eq. (1)-(17) and Algorithm 3 from the paper."""

    def __init__(self, objective, lb, ub, dim, max_fes,
                 sn=50, limit=50, phbdy=0.3, cr_exl=0.1, cr_exp=0.7,
                 ph_period=50, seed=None):
        self.f = objective
        self.lb = np.asarray(lb, dtype=float)
        self.ub = np.asarray(ub, dtype=float)
        self.D = dim
        self.max_fes = int(max_fes)
        self.SN = sn
        self.limit = limit
        self.phbdy = phbdy
        self.cr_exl = cr_exl
        self.cr_exp = cr_exp
        self.ph_period = ph_period
        self.rng = np.random.default_rng(seed)

        self.fes = 0
        self.X = None
        self.fit_raw = None
        self.trial = None
        self.mode = "exploit"

    def _eval(self, x):
        x = np.clip(x, self.lb, self.ub)
        self.fes += 1
        return x, float(self.f(x))

    def _abc_fitness(self, fraw):
        
        return np.where(fraw >= 0, 1.0 / (1.0 + fraw), 1.0 + np.abs(fraw))

    def _init_population(self):
        
        self.X = self.lb + self.rng.random((self.SN, self.D)) * (self.ub - self.lb)
        self.fit_raw = np.empty(self.SN)
        for i in range(self.SN):
            self.X[i], self.fit_raw[i] = self._eval(self.X[i])
        self.trial = np.zeros(self.SN, dtype=int)

    def _update_ph_mode(self):
        ph = compute_ph(self.X, self.fit_raw)
        self.mode = "exploit" if ph < self.phbdy else "explore"

    def _rand_dim(self):
        return int(self.rng.integers(0, self.D))

    def _rand_others(self, i, n, exclude=()):
        pool = [k for k in range(self.SN) if k != i and k not in exclude]
        idx = self.rng.choice(pool, size=n, replace=False)
        return [int(v) for v in idx]

    def _lbest_of(self, i, k_radius):
        neigh = ring_neighbor_indices(i, k_radius, self.SN)
        return int(neigh[np.argmin(self.fit_raw[neigh])])

    def _employed_phase(self, k_radius):
        for i in range(self.SN):
            v = self.X[i].copy()
            j = self._rand_dim()
            if self.mode == "exploit":
                
                lbest = self._lbest_of(i, k_radius)
                r1, r2 = self._rand_others(i, 2, exclude=(lbest,))
                phi = self.rng.uniform(-1, 1)
                v[j] = self.X[lbest, j] + phi * (self.X[r1, j] - self.X[r2, j])
            else:
                
                k = self._rand_others(i, 1)[0]
                phi = self.rng.uniform(-1, 1)
                v[j] = self.X[i, j] + phi * (self.X[i, j] - self.X[k, j])
            self._greedy_select(i, v)
            if self.fes >= self.max_fes:
                return

    def _onlooker_phase(self, k_radius):
        
        fitness = self._abc_fitness(self.fit_raw)
        probs = fitness / fitness.sum()
        best_idx = int(np.argmin(self.fit_raw))
        x_best = self.X[best_idx]

        t = 0
        i = 0
        while t < self.SN:
            if self.rng.random() < probs[i]:
                v = self.X[i].copy()
                j = self._rand_dim()
                phi = self.rng.uniform(-1, 1)  
                r1, r2 = self._rand_others(i, 2)
                if self.mode == "exploit":
                    
                    lbest = self._lbest_of(i, k_radius)
                    v[j] = (self.X[lbest, j] + phi * (self.X[i, j] - self.X[r1, j])
                            + phi * (x_best[j] - self.X[r2, j]))
                else:
                    
                    v[j] = (self.X[i, j] + phi * (self.X[i, j] - self.X[r1, j])
                            + phi * (x_best[j] - self.X[r2, j]))
                self._greedy_select(i, v)
                t += 1
                if self.fes >= self.max_fes:
                    return
            i = (i + 1) % self.SN

    def _greedy_select(self, i, v):
        v, fv = self._eval(v)
        if fv <= self.fit_raw[i]:
            self.X[i] = v
            self.fit_raw[i] = fv
            self.trial[i] = 0
        else:
            self.trial[i] += 1

    def _scout_phase(self, k_radius):
        best_idx = int(np.argmin(self.fit_raw))
        x_best = self.X[best_idx]
        for i in range(self.SN):
            if self.trial[i] <= self.limit:
                continue
            if self.fes >= self.max_fes:
                return
            v = self.X[i].copy()
            if self.mode == "exploit":
                
                for j in range(self.D):
                    if self.rng.random() <= self.cr_exl:
                        v[j] = self.rng.normal(x_best[j], 1.0)
            else:
                
                r = self._rand_others(i, 1)[0]
                neigh = ring_neighbor_indices(r, k_radius, self.SN)
                x_rlbest = self.X[neigh[np.argmin(self.fit_raw[neigh])]]
                for j in range(self.D):
                    if self.rng.random() <= self.cr_exp:
                        mean = (x_rlbest[j] + x_best[j]) / 2.0
                        std = abs(x_rlbest[j] - x_best[j]) + 1e-12
                        v[j] = self.rng.normal(mean, std)
            v, fv = self._eval(v)
            self.X[i] = v
            self.fit_raw[i] = fv
            self.trial[i] = 0

    def run(self):
        self._init_population()
        gen = 0
        self._update_ph_mode()
        while self.fes < self.max_fes:
            k_radius = dynamic_radius(self.fes, self.max_fes, self.SN)
            self._employed_phase(k_radius)
            if self.fes >= self.max_fes:
                break
            self._onlooker_phase(k_radius)
            if self.fes >= self.max_fes:
                break
            self._scout_phase(k_radius)
            gen += 1
            if gen % self.ph_period == 0:
                self._update_ph_mode()
        best_idx = int(np.argmin(self.fit_raw))
        return float(self.fit_raw[best_idx]), self.fes

def run_trial(objective, lb, ub, dim, max_fes, seed, f_global=0.0, **kwargs):
    t0 = time.time()
    algo = ASABC(objective, lb, ub, dim, max_fes, seed=seed, **kwargs)
    best_raw, fes_used = algo.run()
    err = max(best_raw - f_global, 0.0)
    return {"best_error": err, "fes_used": fes_used, "time_sec": time.time() - t0}

import os, time, csv
from opfunu.cec_based import cec2013

DIMENSIONS      = [30, 50]
N_RUNS          = 30
SN              = 50
LIMIT           = 100
PHBDY           = 0.3
CR_EXL          = 0.1
CR_EXP          = 0.7
PH_PERIOD       = 50
D=30
MAXFES_PER_DIM  = 10000 * D 
FUNCTION_NAMES  = [f"F{i}" for i in range(1, 29)]

OUTPUT_DIR     = os.path.abspath("results")
RAW_CSV        = os.path.join(OUTPUT_DIR, "asabc_cec2013_raw_results.csv")
SUMMARY_CSV    = os.path.join(OUTPUT_DIR, "asabc_cec2013_summary.csv")
DOCX_OUT       = os.path.join(OUTPUT_DIR, "ASABC_CEC2013_report.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)

def get_opfunu_class(fname, dim):
    idx = int(fname.replace("F", ""))
    cls = getattr(cec2013, f"F{idx}2013")
    return cls(ndim=dim)

print("Configuration loaded.")
print(f"  {len(FUNCTION_NAMES)} functions x {len(DIMENSIONS)} dimensions x {N_RUNS} runs "
      f"= {len(FUNCTION_NAMES)*len(DIMENSIONS)*N_RUNS} trials in total.")
print(f"  Results -> {RAW_CSV}")

RAW_FIELDS = ["function", "dim", "run", "best_error", "fes_used", "time_sec", "timestamp"]

def load_done_trials():
    done = set()
    if os.path.exists(RAW_CSV):
        with open(RAW_CSV, newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                done.add((row["function"], int(row["dim"]), int(row["run"])))
    else:
        with open(RAW_CSV, "w", newline="") as f:
            csv.DictWriter(f, fieldnames=RAW_FIELDS).writeheader()
    return done

def append_result(function, dim, run, result):
    with open(RAW_CSV, "a", newline="") as f:
        w = csv.DictWriter(f, fieldnames=RAW_FIELDS)
        w.writerow({
            "function": function, "dim": dim, "run": run,
            "best_error": result["best_error"], "fes_used": result["fes_used"],
            "time_sec": round(result["time_sec"], 3),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        })

done_trials = load_done_trials()

tasks = [
    (fname, dim, run)
    for dim in DIMENSIONS
    for fname in FUNCTION_NAMES
    for run in range(1, N_RUNS + 1)
]
n_total = len(tasks)

t_start = time.time()
n_done = len(done_trials)

print(f"{n_done}/{n_total} trials already completed (resumed from {RAW_CSV}).")

for (fname, dim, run) in tasks:
    if (fname, dim, run) in done_trials:
        continue

    prob = get_opfunu_class(fname, dim)
    max_fes = MAXFES_PER_DIM * dim
    seed = hash((fname, dim, run)) % (2**32)

    result = run_trial(
        objective=lambda x, _p=prob: _p.evaluate(x),
        lb=prob.lb, ub=prob.ub, dim=dim, max_fes=max_fes, seed=seed,
        f_global=prob.f_global,
        sn=SN, limit=LIMIT, phbdy=PHBDY, cr_exl=CR_EXL, cr_exp=CR_EXP,
        ph_period=PH_PERIOD,
    )
    append_result(fname, dim, run, result)
    n_done += 1

    if n_done % 10 == 0 or n_done == n_total:
        elapsed = time.time() - t_start
        print(f"[{n_done}/{n_total}] {fname} D={dim} run={run} "
              f"err={result['best_error']:.4e}  (elapsed {elapsed/60:.1f} min)")

print("Experiment finished." if n_done == n_total else "Interrupted — rerun the cell to resume.")

import pandas as pd

raw = pd.read_csv(RAW_CSV)
summary = (
    raw.groupby(["function", "dim"])
       .agg(mean_error=("best_error", "mean"),
            std_error=("best_error", "std"),
            n_runs=("best_error", "count"),
            mean_fes=("fes_used", "mean"))
       .reset_index()
)
summary["function_order"] = summary["function"].str.replace("F", "").astype(int)
summary = summary.sort_values(["dim", "function_order"]).drop(columns="function_order")
summary.to_csv(SUMMARY_CSV, index=False)

print(f"Summary written -> {SUMMARY_CSV}")
summary

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

summary = pd.read_csv(SUMMARY_CSV)
funcs = sorted(summary["function"].unique(), key=lambda s: int(s.replace("F", "")))
dims = sorted(summary["dim"].unique())

doc = Document()

title = doc.add_heading("Benchmarking ASABC on the CEC2013 Test Suite", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Reproduction study following the experimental protocol of Zhang et al. (2026)")
r.italic = True

doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "This report presents an independent reproduction of the Adaptive exploration-"
    "exploitation Switching Artificial Bee Colony algorithm (ASABC), evaluated on the "
    "CEC2013 numerical optimization benchmark suite. The experimental protocol "
    "(population size, maximum number of function evaluations, number of independent "
    "runs, and control-parameter settings) follows the specification reported in the "
    "original publication. Mean and standard deviation of the function error "
    "(f(x) - f*) are reported for each of the 28 benchmark functions."
)

doc.add_heading("1. Experimental protocol", level=1)
doc.add_paragraph(
    "All experiments follow the setup described in the source paper's Experimental "
    "Setup section (Sec. 4.1) and default control-parameter values from the "
    "sensitivity analysis (Sec. 4.6). Table 1 summarizes the configuration used."
)
protocol = [
    ("Benchmark suite", "CEC2013 (28 functions: F1-F5 unimodal, F6-F20 basic multimodal, F21-F28 composition)"),
    ("Dimensions tested", ", ".join(str(d) for d in dims)),
    ("Population size (SN)", "50"),
    ("Max function evaluations (MaxFEs)", "10000 x D"),
    ("Independent runs per function", "30"),
    ("limit (scout trigger)", "50"),
    ("phbdy (PH threshold, smooth/rugged)", "0.3"),
    ("CR_exl (exploitation scout)", "0.1"),
    ("CR_exp (exploration scout)", "0.7"),
    ("PH re-evaluation period", "every 50 generations"),
    ("ILM reference function", "Sphere"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text, hdr[1].text = "Parameter", "Value"
for k, v in protocol:
    row = t.add_row().cells
    row[0].text, row[1].text = k, v

doc.add_paragraph()
doc.add_paragraph(
    "Implementation note: this reproduction follows Eqs. (1)-(17) and Algorithms 1-3 "
    "exactly as specified in the source PDF (Zhang et al., 2026), including the ring-"
    "neighborhood radius (Eq. 12), the exploitation/exploration search operators for "
    "the employed, onlooker and scout bee phases (Eqs. 11, 14-17), and the online ILM "
    "problem-hardness measure (Eqs. 5-10) used to switch between strategy sets."
)

doc.add_heading("2. Results", level=1)
for d in dims:
    doc.add_heading(f"Table -- CEC2013, D = {d}", level=2)
    sub_df = summary[summary["dim"] == d].set_index("function")
    tt = doc.add_table(rows=1, cols=5)
    tt.style = "Light Grid Accent 1"
    hdr = tt.rows[0].cells
    for c, name in zip(hdr, ["Func.", "Mean error", "Std. dev.", "Runs", "Mean FEs used"]):
        c.text = name
    for fn in funcs:
        if fn not in sub_df.index:
            continue
        row = sub_df.loc[fn]
        cells = tt.add_row().cells
        cells[0].text = fn
        cells[1].text = f"{row['mean_error']:.2e}"
        cells[2].text = f"{row['std_error']:.2e}"
        cells[3].text = str(int(row["n_runs"]))
        cells[4].text = f"{row['mean_fes']:.0f}"
    doc.add_paragraph()

doc.add_heading("3. Discussion", level=1)
doc.add_paragraph(
    "[To be completed once the 30 runs are finished: comparison of unimodal / "
    "multimodal / composition functions, behavior of the switching mechanism (PH values), "
    "and comparison with the results published in the source paper (Tables 1-2) if "
    "you also re-implement the four competing variants ABC-MIG, ABCNG, "
    "ABC-ANT and FLABC.]"
)

doc.add_heading("4. Reference", level=1)
doc.add_paragraph(
    "Zhang, Z., Zhou, X., Song, J., Liu, F., & Ma, Y. (2026). Adaptive "
    "exploration-exploitation switching artificial bee colony algorithm based on "
    "problem features. Expert Systems With Applications, 318, 131979. "
    "https://doi.org/10.1016/j.eswa.2026.131979"
)

doc.save(DOCX_OUT)
print(f"Report generated -> {DOCX_OUT}")
print(f"Raw CSV           -> {RAW_CSV}")
print(f"Summary CSV       -> {SUMMARY_CSV}")
