#!/usr/bin/env python
# coding: utf-8

# In[3]:


print("Ready.")

# In[4]:


import numpy as np
import pandas as pd
from pathlib import Path
from tqdm.auto import tqdm

from opfunu.cec_based import cec2013

RNG_MASTER_SEED = 42
OUTPUT_DIR = Path("results/NABC_D50")
OUTPUT_DIR.mkdir(exist_ok=True)

print("Setup OK.")

# In[5]:


def fitness_transform(f):
    return np.where(f >= 0, 1.0 / (1.0 + f), 1.0 + np.abs(f))


def nabc(evaluate, lb, ub, D, SN, limit, N, max_nfe, rng):
    lb = np.asarray(lb, dtype=float)
    ub = np.asarray(ub, dtype=float)

    X = lb + rng.random((SN, D)) * (ub - lb)
    fit = np.empty(SN)
    for i in range(SN):
        fit[i] = evaluate(X[i])
    nfe = SN
    trial = np.zeros(SN, dtype=int)

    all_idx = np.arange(SN)

    def best_neighbor_move(i, X, fit):
        candidates = rng.choice(SN - 1, size=min(N, SN - 1), replace=False)
        candidates = np.where(candidates >= i, candidates + 1, candidates)
        nbest = candidates[np.argmin(fit[candidates])]
        j = rng.integers(0, D)
        phi = rng.uniform(-1.0, 1.0)
        V = X[i].copy()
        V[j] = X[nbest, j] + phi * (X[nbest, j] - X[i, j])
        np.clip(V, lb, ub, out=V)
        return V

    while nfe < max_nfe:
        for i in range(SN):
            if nfe >= max_nfe:
                break
            V = best_neighbor_move(i, X, fit)
            fV = evaluate(V)
            nfe += 1
            if fit[i] > fV:
                X[i] = V
                fit[i] = fV
                trial[i] = 0
            else:
                trial[i] += 1

        if nfe >= max_nfe:
            break

        fitv = fitness_transform(fit)
        p = fitv / fitv.sum()
        count = 0
        while count < SN and nfe < max_nfe:
            j = rng.choice(SN, p=p)
            V = best_neighbor_move(j, X, fit)
            fV = evaluate(V)
            nfe += 1
            if fit[j] > fV:
                X[j] = V
                fit[j] = fV
                trial[j] = 0
            else:
                trial[j] += 1
            count += 1

        if nfe >= max_nfe:
            break

        gbest_idx = np.argmin(fit)
        for i in range(SN):
            if nfe >= max_nfe:
                break
            if trial[i] > limit:
                trial[i] = 0
                others = all_idx[all_idx != i]
                a, b = rng.choice(others, size=2, replace=False)
                r = rng.random(3)
                r = r / r.sum()
                V = r[0] * X[i] + r[1] * X[gbest_idx] + r[2] * (X[a] - X[b])
                np.clip(V, lb, ub, out=V)
                fV = evaluate(V)
                nfe += 1
                X[i] = V
                fit[i] = fV

    return float(np.min(fit))

# In[ ]:


D = 50
SN = 50
LIMIT = 100
NEIGH_N = 5
MAX_NFE = 10000 * D
N_RUNS = 30

def make_functions(ndim=D):
    funcs = {}
    for i in range(1, 29):
        cls = getattr(cec2013, f"F{i}2013")
        inst = cls(ndim=ndim)
        funcs[f"f{i}"] = inst
    return funcs

FUNCTIONS = make_functions(D)
print(f"{len(FUNCTIONS)} CEC2013 functions loaded (D={D}).")
for k, f in list(FUNCTIONS.items())[:3]:
    print(k, "->", f.name, "| f_global =", f.f_global)

# In[7]:


def run_single(func_name, run_idx):
    f_inst = getattr(cec2013, f"F{func_name[1:]}2013")(ndim=D)
    seed = RNG_MASTER_SEED * 10_000 + int(func_name[1:]) * 100 + run_idx
    rng = np.random.default_rng(seed)
    best = nabc(f_inst.evaluate, f_inst.lb, f_inst.ub, D, SN, LIMIT, NEIGH_N, MAX_NFE, rng)
    error = best - f_inst.f_global
    return max(error, 0.0)


def run_experiment():
    results = {}

    func_names = [f"f{i}" for i in range(1, 29)]
    pbar = tqdm(func_names, desc="CEC2013 functions")
    for func_name in pbar:
        pbar.set_postfix(function=func_name)
        errors = [run_single(func_name, run_idx) for run_idx in range(N_RUNS)]
        results[func_name] = errors

    return results

# In[8]:


results = run_experiment()
print("Done.")

# In[9]:


rows = []
for i in range(1, 29):
    fname = f"f{i}"
    errors = np.array(results[fname])
    f_inst = FUNCTIONS[fname]
    rows.append({
        "Function": f"f{i}",
        "Name": f_inst.name,
        "D": D,
        "Runs": len(errors),
        "Mean Error": errors.mean(),
        "Std Dev": errors.std(ddof=1) if len(errors) > 1 else 0.0,
        "Best": errors.min(),
        "Worst": errors.max(),
    })

df_results = pd.DataFrame(rows)
pd.set_option("display.float_format", lambda x: f"{x:.4E}")
df_results

# In[ ]:


csv_path = OUTPUT_DIR / f"nabc_cec2013_D{D}_results.csv"
df_out = df_results.copy()
for col in ["Mean Error", "Std Dev", "Best", "Worst"]:
    df_out[col] = df_out[col].map(lambda x: f"{x:.6E}")
df_out.to_csv(csv_path, index=False)
print(f"CSV saved: {csv_path.resolve()}")
df_out

# In[11]:


from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_report(df, path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)

    title = doc.add_heading(
        "Reproduction of NABC on the CEC 2013 Benchmark Suite", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Reproduced from: Peng, H., Deng, C., & Wu, Z. (2018). Best neighbor-guided "
        "artificial bee colony algorithm for continuous optimization problems. "
        "Soft Computing. https://doi.org/10.1007/s00500-018-3473-6"
    ).italic = True

    doc.add_heading("Experimental Protocol", level=2)
    proto = doc.add_paragraph()
    proto.add_run(
        f"Dimension D = {D}; Population size SN = D = {SN}; limit = {LIMIT}; "
        f"neighbor size N = {NEIGH_N}; MaxNFE = 5000 x D = {MAX_NFE}; "
        f"{N_RUNS} independent runs per function. Error = f(x*) - f_optimal, "
        f"reported as Mean +/- Std over 30 runs, following Section 5.7 of the original paper."
    )

    doc.add_heading("Results on CEC 2013 (D = 30)", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Function", "Optimum", "Mean Error", "Std Dev"]
    for cell, text in zip(hdr, headers):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9D9D9")

    for _, row in df.iterrows():
        fname = row["Function"]
        f_inst = FUNCTIONS[fname]
        cells = table.add_row().cells
        cells[0].text = fname
        cells[1].text = f"{f_inst.f_global:.1f}"
        cells[2].text = f"{row['Mean Error']:.4E}"
        cells[3].text = f"{row['Std Dev']:.4E}"
        for c in cells:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Note: results were generated with an independent re-implementation of NABC "
        "based on Algorithm 1 and Equations (1)-(5) of the source paper, using the "
        "opfunu CEC2013 benchmark suite (30 independent runs per function)."
    ).italic = True

    doc.save(path)


docx_path = OUTPUT_DIR / "NABC_CEC2013_Report.docx"
build_report(df_results, docx_path)
print(f"DOCX saved: {docx_path.resolve()}")
